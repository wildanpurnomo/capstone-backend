import io
import requests
import docx
import sys
from nltk.tokenize import RegexpTokenizer
from docx import Document
from pathlib import Path
from docx import Document

def processDocx(inputUrl):
    copyPath = Path().absolute().joinpath('test.docx')

    document = requests.get(inputUrl, allow_redirects=True)
    document = Document(io.BytesIO(document.content))
    document.save(copyPath)
    document = docx.Document(copyPath)
    docText = '\n'.join(
        paragraph.text for paragraph in document.paragraphs
    )
    ''.join(docText)
    texts = docText.split('\n')
    texts = list(filter(None, texts))

    result = ''
    for i in range(len(texts)):
        tokenizer = RegexpTokenizer(r'\w+')
        textsTokenized = tokenizer.tokenize(texts[i])
        s = " "
        texts[i] = s.join(textsTokenized).lstrip(' ')
        result = s.join(texts)

    return result

def main():
    print(processDocx(sys.argv[1]))

if __name__ == '__main__':
    main()
